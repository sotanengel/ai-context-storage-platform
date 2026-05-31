"""Silver converter: DOCX → CDM via python-docx."""

from __future__ import annotations

import io
from typing import Any

from formaforge.models.silver import (
    CdmBlock,
    CdmDocument,
    CdmFrontmatter,
    CdmTableBlock,
    ConversionMethod,
)
from formaforge.silver.converters.base import BaseConverter


def _open_docx(data: bytes) -> Any:
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError(
            "python-docx is required for DOCX ingestion. Install with: uv add 'formaforge[binary]'"
        ) from exc
    return Document(io.BytesIO(data))


class DocxConverter(BaseConverter):
    source_format = "docx"

    def convert(self, raw: str, source_uri: str = "") -> CdmDocument:
        return self.convert_bytes(raw.encode("utf-8", errors="replace"), source_uri)

    def convert_bytes(self, raw: bytes, source_uri: str = "") -> CdmDocument:
        word_doc = _open_docx(raw)
        title = ""
        body_parts: list[str] = []

        for para in word_doc.paragraphs:
            if not para.text.strip():
                continue
            if para.style.name.startswith("Heading 1") and not title:
                title = para.text
            else:
                body_parts.append(para.text)

        blocks: list[CdmBlock] = []
        for table in word_doc.tables:
            rows = table.rows
            if not rows:
                continue
            headers = [cell.text.strip() for cell in rows[0].cells]
            columns = [f"{h}:string" for h in headers]
            data_rows = [[cell.text.strip() for cell in row.cells] for row in rows[1:]]
            blocks.append(CdmTableBlock(columns=columns, rows=data_rows))

        fm = CdmFrontmatter(
            source_format=self.source_format,
            source_uri=source_uri,
            structure_class="unstructured",
            conversion_method=ConversionMethod.DETERMINISTIC,
            conversion_confidence=0.95,
        )
        return CdmDocument(frontmatter=fm, title=title, body="\n".join(body_parts), blocks=blocks)
